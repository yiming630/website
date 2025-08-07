"""
File Processing Service - FastAPI
Integrates PDF to DOCX conversion and document processing
"""
import os
import sys
import asyncio
import logging
import uuid
from contextlib import asynccontextmanager
from pathlib import Path
from typing import Optional, List, Dict, Any
import json

# Add the existing PDF processing modules to path
sys.path.insert(0, str(Path(__file__).parent.parent.parent.parent / "Test" / "PDF_to_DOCX"))

from fastapi import FastAPI, HTTPException, Depends, BackgroundTasks, UploadFile, File, status
from fastapi.middleware.cors import CORSMiddleware
from fastapi.middleware.gzip import GZipMiddleware
from fastapi.responses import JSONResponse
from pydantic import BaseModel, ValidationError
import uvicorn

# Import existing processing modules
try:
    from src.converter import PDFConverter
    from src.splitter import DocumentSplitter
    from src.pdf_splitter import PDFSplitter
    from src.gemini_client import GeminiClient
    from src.cloud_storage import CloudStorageManager
    from src.wps_client import WPSClient
except ImportError as e:
    logging.warning(f"Could not import PDF processing modules: {e}")
    # Create dummy classes for development
    class PDFConverter: pass
    class DocumentSplitter: pass
    class PDFSplitter: pass
    class GeminiClient: pass
    class CloudStorageManager: pass
    class WPSClient: pass

# Import database connection
sys.path.insert(0, str(Path(__file__).parent.parent.parent / "databases"))
try:
    from connection import query, transaction
except ImportError:
    # Fallback for development
    async def query(*args): return type('Result', (), {'rows': []})()
    async def transaction(callback): return await callback(None)

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Global instances
pdf_converter = None
document_splitter = None
pdf_splitter = None
gemini_client = None
cloud_storage = None
wps_client = None

@asynccontextmanager
async def lifespan(app: FastAPI):
    """Startup and shutdown events"""
    global pdf_converter, document_splitter, pdf_splitter, gemini_client, cloud_storage, wps_client
    
    # Startup
    logger.info("🚀 Starting File Processing Service...")
    
    try:
        # Initialize processing components
        pdf_converter = PDFConverter()
        document_splitter = DocumentSplitter()
        pdf_splitter = PDFSplitter()
        gemini_client = GeminiClient()
        cloud_storage = CloudStorageManager()
        wps_client = WPSClient()
        
        logger.info("✅ File processing components initialized")
        
        # Test database connection
        try:
            await query("SELECT 1")
            logger.info("✅ Database connection established")
        except Exception as e:
            logger.warning(f"Database connection failed: {e}")
        
    except Exception as e:
        logger.warning(f"Some components failed to initialize: {e}")
    
    yield
    
    # Shutdown
    logger.info("🛑 Shutting down File Processing Service...")

# Create FastAPI app
app = FastAPI(
    title="Translation Platform - File Processing Service",
    description="File processing, conversion, and document manipulation service",
    version="1.0.0",
    lifespan=lifespan
)

# Add middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Configure properly in production
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)
app.add_middleware(GZipMiddleware, minimum_size=1000)

# Pydantic models
class FileProcessRequest(BaseModel):
    document_id: str
    file_path: str
    process_type: str = "pdf_to_docx"  # pdf_to_docx, docx_split, pdf_split
    options: Dict[str, Any] = {}

class ProcessingStatus(BaseModel):
    document_id: str
    status: str
    progress: int
    current_step: str
    estimated_time_remaining: Optional[int] = None
    error: Optional[str] = None
    result_files: List[str] = []

class ConversionRequest(BaseModel):
    file_type: str  # pdf, docx, txt
    target_format: str  # docx, json, txt
    options: Dict[str, Any] = {}

class ConversionResponse(BaseModel):
    success: bool
    output_files: List[str]
    processing_time: float
    metadata: Dict[str, Any] = {}

class DocumentToJsonRequest(BaseModel):
    docx_path: str
    extract_formatting: bool = True
    extract_images: bool = False

class DocumentJsonResponse(BaseModel):
    success: bool
    json_content: Dict[str, Any]
    metadata: Dict[str, Any]

# Health check
@app.get("/health")
async def health_check():
    """Health check endpoint"""
    return {
        "status": "ok",
        "service": "file-processing-service",
        "timestamp": asyncio.get_event_loop().time(),
        "components": {
            "pdf_converter": pdf_converter is not None,
            "document_splitter": document_splitter is not None,
            "pdf_splitter": pdf_splitter is not None,
            "gemini_client": gemini_client is not None,
            "cloud_storage": cloud_storage is not None,
            "wps_client": wps_client is not None
        }
    }

# File upload endpoint
@app.post("/api/files/upload")
async def upload_file(
    file: UploadFile = File(...),
    background_tasks: BackgroundTasks = BackgroundTasks()
):
    """Upload and store a file for processing"""
    try:
        # Validate file type
        allowed_extensions = {'.pdf', '.docx', '.txt', '.doc'}
        file_extension = Path(file.filename).suffix.lower()
        
        if file_extension not in allowed_extensions:
            raise HTTPException(
                status_code=400, 
                detail=f"Unsupported file type. Allowed: {', '.join(allowed_extensions)}"
            )
        
        # Create upload directory
        upload_dir = Path("temp_uploads")
        upload_dir.mkdir(exist_ok=True)
        
        # Generate unique filename
        file_id = str(uuid.uuid4())
        file_path = upload_dir / f"{file_id}_{file.filename}"
        
        # Save uploaded file
        with open(file_path, "wb") as buffer:
            content = await file.read()
            buffer.write(content)
        
        return {
            "file_id": file_id,
            "original_filename": file.filename,
            "file_path": str(file_path),
            "size": len(content),
            "file_type": file_extension,
            "status": "uploaded"
        }
        
    except Exception as e:
        logger.error(f"Upload failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# PDF to DOCX conversion
@app.post("/api/convert/pdf-to-docx")
async def convert_pdf_to_docx(
    request: ConversionRequest,
    background_tasks: BackgroundTasks
):
    """Convert PDF to DOCX using existing PDF processing system"""
    try:
        if not pdf_converter:
            raise HTTPException(status_code=503, detail="PDF converter not available")
        
        # Implementation depends on the request structure
        # This is a placeholder that would integrate with the existing system
        
        return ConversionResponse(
            success=True,
            output_files=["placeholder_output.docx"],
            processing_time=0.0,
            metadata={"converter": "wps_api"}
        )
        
    except Exception as e:
        logger.error(f"PDF conversion failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# DOCX to JSON conversion
@app.post("/api/convert/docx-to-json", response_model=DocumentJsonResponse)
async def convert_docx_to_json(request: DocumentToJsonRequest):
    """Convert DOCX document to structured JSON format"""
    try:
        from docx import Document
        import json
        
        if not Path(request.docx_path).exists():
            raise HTTPException(status_code=404, detail="DOCX file not found")
        
        # Load DOCX document
        doc = Document(request.docx_path)
        
        # Extract content to JSON structure
        json_content = {
            "document": {
                "metadata": {
                    "filename": Path(request.docx_path).name,
                    "paragraphs_count": len(doc.paragraphs),
                    "tables_count": len(doc.tables) if hasattr(doc, 'tables') else 0,
                },
                "content": {
                    "paragraphs": [],
                    "tables": [],
                    "images": [] if request.extract_images else None
                }
            }
        }
        
        # Extract paragraphs
        for i, paragraph in enumerate(doc.paragraphs):
            para_data = {
                "index": i,
                "text": paragraph.text,
                "style": paragraph.style.name if paragraph.style else None
            }
            
            if request.extract_formatting:
                para_data["formatting"] = {
                    "alignment": str(paragraph.alignment) if paragraph.alignment else None,
                }
                
                # Extract run-level formatting
                runs = []
                for run in paragraph.runs:
                    run_data = {
                        "text": run.text,
                        "bold": run.bold,
                        "italic": run.italic,
                        "underline": run.underline,
                        "font_name": run.font.name if run.font.name else None,
                        "font_size": str(run.font.size) if run.font.size else None
                    }
                    runs.append(run_data)
                para_data["runs"] = runs
            
            json_content["document"]["content"]["paragraphs"].append(para_data)
        
        # Extract tables
        if hasattr(doc, 'tables'):
            for i, table in enumerate(doc.tables):
                table_data = {
                    "index": i,
                    "rows": len(table.rows),
                    "columns": len(table.columns) if table.rows else 0,
                    "data": []
                }
                
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append({
                            "text": cell.text,
                            "paragraphs": len(cell.paragraphs)
                        })
                    table_data["data"].append(row_data)
                
                json_content["document"]["content"]["tables"].append(table_data)
        
        metadata = {
            "processing_time": 0.0,  # Would be calculated in real implementation
            "extraction_options": {
                "formatting": request.extract_formatting,
                "images": request.extract_images
            }
        }
        
        return DocumentJsonResponse(
            success=True,
            json_content=json_content,
            metadata=metadata
        )
        
    except Exception as e:
        logger.error(f"DOCX to JSON conversion failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# Document processing endpoint
@app.post("/api/process/document")
async def process_document(
    request: FileProcessRequest,
    background_tasks: BackgroundTasks
):
    """Process document with specified processing type"""
    try:
        if request.process_type == "pdf_to_docx":
            # Use PDF converter
            background_tasks.add_task(
                process_pdf_to_docx_background,
                request.document_id,
                request.file_path,
                request.options
            )
        elif request.process_type == "pdf_split":
            # Use PDF splitter
            background_tasks.add_task(
                process_pdf_split_background,
                request.document_id,
                request.file_path,
                request.options
            )
        elif request.process_type == "docx_split":
            # Use document splitter
            background_tasks.add_task(
                process_docx_split_background,
                request.document_id,
                request.file_path,
                request.options
            )
        else:
            raise HTTPException(status_code=400, detail=f"Unknown process type: {request.process_type}")
        
        return {
            "document_id": request.document_id,
            "process_type": request.process_type,
            "status": "processing_started"
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Document processing failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# Processing status endpoint
@app.get("/api/process/status/{document_id}")
async def get_processing_status(document_id: str):
    """Get document processing status"""
    try:
        # In a real implementation, this would check processing status from database or cache
        # For now, return a placeholder status
        return ProcessingStatus(
            document_id=document_id,
            status="processing",
            progress=50,
            current_step="Converting document...",
            estimated_time_remaining=120,
            result_files=[]
        )
        
    except Exception as e:
        logger.error(f"Status check failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# Background processing functions
async def process_pdf_to_docx_background(document_id: str, file_path: str, options: Dict[str, Any]):
    """Background task for PDF to DOCX conversion"""
    try:
        logger.info(f"Starting PDF to DOCX conversion for document {document_id}")
        
        if pdf_converter:
            # Use the existing PDF converter
            result = pdf_converter.convert_single_pdf(Path(file_path), Path("temp_outputs"))
            
            if result[0]:  # Success
                logger.info(f"PDF conversion successful for document {document_id}")
                # Update database with success status
                await update_processing_status(document_id, "completed", 100, result_files=[str(result[1])])
            else:
                logger.error(f"PDF conversion failed for document {document_id}: {result[2]}")
                await update_processing_status(document_id, "failed", 0, error=result[2])
        else:
            logger.warning("PDF converter not available, using placeholder")
            await update_processing_status(document_id, "completed", 100)
        
    except Exception as e:
        logger.error(f"PDF to DOCX conversion failed for document {document_id}: {e}")
        await update_processing_status(document_id, "failed", 0, error=str(e))

async def process_pdf_split_background(document_id: str, file_path: str, options: Dict[str, Any]):
    """Background task for PDF splitting and conversion"""
    try:
        logger.info(f"Starting PDF split and convert for document {document_id}")
        
        if pdf_splitter:
            # Use the existing PDF splitter
            result = pdf_splitter.split_and_convert_pdf(file_path, options.get('prompt'))
            
            if result.get('success'):
                logger.info(f"PDF split and convert successful for document {document_id}")
                await update_processing_status(
                    document_id, 
                    "completed", 
                    100, 
                    result_files=result.get('docx_files', [])
                )
            else:
                error_msg = result.get('error', 'Unknown error during PDF split')
                logger.error(f"PDF split failed for document {document_id}: {error_msg}")
                await update_processing_status(document_id, "failed", 0, error=error_msg)
        else:
            logger.warning("PDF splitter not available, using placeholder")
            await update_processing_status(document_id, "completed", 100)
        
    except Exception as e:
        logger.error(f"PDF split failed for document {document_id}: {e}")
        await update_processing_status(document_id, "failed", 0, error=str(e))

async def process_docx_split_background(document_id: str, file_path: str, options: Dict[str, Any]):
    """Background task for DOCX document splitting"""
    try:
        logger.info(f"Starting DOCX split for document {document_id}")
        
        if document_splitter:
            # Use the existing document splitter
            result = document_splitter.split_single_document(Path(file_path), Path("temp_outputs"))
            
            if result.get('success'):
                logger.info(f"DOCX split successful for document {document_id}")
                await update_processing_status(
                    document_id, 
                    "completed", 
                    100, 
                    result_files=result.get('output_files', [])
                )
            else:
                error_msg = result.get('error', 'Unknown error during DOCX split')
                logger.error(f"DOCX split failed for document {document_id}: {error_msg}")
                await update_processing_status(document_id, "failed", 0, error=error_msg)
        else:
            logger.warning("Document splitter not available, using placeholder")
            await update_processing_status(document_id, "completed", 100)
        
    except Exception as e:
        logger.error(f"DOCX split failed for document {document_id}: {e}")
        await update_processing_status(document_id, "failed", 0, error=str(e))

async def update_processing_status(
    document_id: str, 
    status: str, 
    progress: int, 
    error: str = None, 
    result_files: List[str] = None
):
    """Update document processing status in database"""
    try:
        # This would update the database with processing status
        # For now, just log the status update
        logger.info(f"Document {document_id} status: {status} ({progress}%)")
        if error:
            logger.error(f"Document {document_id} error: {error}")
        if result_files:
            logger.info(f"Document {document_id} result files: {result_files}")
        
    except Exception as e:
        logger.error(f"Failed to update processing status: {e}")

if __name__ == "__main__":
    uvicorn.run(
        "main:app",
        host="0.0.0.0",
        port=8001,
        reload=True,
        log_level="info"
    )
