#!/usr/bin/env python3
"""
多格式导出器模块
将翻译后的DOCX母版转换为用户指定的各种格式
"""

import pathlib
import subprocess
import platform
import tempfile
import shutil
from typing import Optional, List, Dict, Callable

class FormatExporter:
    """格式导出器基类"""
    
    def __init__(self):
        self.system = platform.system()
        
    def export_to_pdf(self, docx_path: pathlib.Path, output_path: Optional[pathlib.Path] = None) -> pathlib.Path:
        """导出为PDF格式"""
        if output_path is None:
            output_path = docx_path.with_suffix('.pdf')
            
        print(f"   📄 导出PDF: {output_path.name}")
        
        try:
            # 优先使用docx2pdf（Windows/macOS下效果最好）
            from docx2pdf import convert
            convert(str(docx_path), str(output_path))
            print(f"   ✅ PDF导出成功（使用docx2pdf）")
            return output_path
        except ImportError:
            print(f"   ⚠️  docx2pdf未安装，尝试使用LibreOffice...")
            
        # 回退方案：使用LibreOffice
        try:
            if self.system == "Windows":
                # Windows下LibreOffice路径
                soffice_paths = [
                    r"C:\Program Files\LibreOffice\program\soffice.exe",
                    r"C:\Program Files (x86)\LibreOffice\program\soffice.exe"
                ]
                soffice = None
                for path in soffice_paths:
                    if pathlib.Path(path).exists():
                        soffice = path
                        break
                
                if not soffice:
                    raise FileNotFoundError("未找到LibreOffice安装")
                    
                cmd = [soffice, "--headless", "--convert-to", "pdf", 
                       "--outdir", str(docx_path.parent), str(docx_path)]
            else:
                # Linux/macOS
                cmd = ["soffice", "--headless", "--convert-to", "pdf",
                       "--outdir", str(docx_path.parent), str(docx_path)]
                       
            subprocess.run(cmd, check=True, capture_output=True)
            print(f"   ✅ PDF导出成功（使用LibreOffice）")
            return output_path
            
        except Exception as e:
            print(f"   ❌ PDF导出失败: {e}")
            raise
            
    def export_to_html(self, docx_path: pathlib.Path, output_path: Optional[pathlib.Path] = None) -> pathlib.Path:
        """导出为HTML格式"""
        if output_path is None:
            output_path = docx_path.with_suffix('.html')
            
        print(f"   🌐 导出HTML: {output_path.name}")
        
        try:
            import mammoth
            
            # 自定义样式映射，保留更多格式
            style_map = """
            p[style-name='Heading 1'] => h1:fresh
            p[style-name='Heading 2'] => h2:fresh
            p[style-name='Heading 3'] => h3:fresh
            b => b
            i => i
            u => u
            """
            
            with open(docx_path, "rb") as docx_file:
                result = mammoth.convert_to_html(
                    docx_file,
                    style_map=style_map,
                    include_embedded_style_map=False
                )
                
            # 添加基础CSS样式
            html_content = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{docx_path.stem}</title>
    <style>
        body {{
            font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, "Helvetica Neue", Arial, sans-serif;
            line-height: 1.6;
            max-width: 800px;
            margin: 0 auto;
            padding: 20px;
            color: #333;
        }}
        h1, h2, h3 {{ color: #2c3e50; }}
        table {{
            border-collapse: collapse;
            width: 100%;
            margin: 20px 0;
        }}
        th, td {{
            border: 1px solid #ddd;
            padding: 8px;
            text-align: left;
        }}
        th {{
            background-color: #f2f2f2;
            font-weight: bold;
        }}
        img {{
            max-width: 100%;
            height: auto;
        }}
    </style>
</head>
<body>
{result.value}
</body>
</html>"""
            
            output_path.write_text(html_content, encoding='utf-8')
            
            # 保存转换消息（如果有）
            if result.messages:
                messages_path = output_path.with_suffix('.html.log')
                with open(messages_path, 'w', encoding='utf-8') as f:
                    for message in result.messages:
                        f.write(f"{message}\n")
                        
            print(f"   ✅ HTML导出成功")
            return output_path
            
        except ImportError:
            print(f"   ❌ 需要安装mammoth: pip install mammoth")
            raise
        except Exception as e:
            print(f"   ❌ HTML导出失败: {e}")
            raise
            
    def export_to_markdown(self, docx_path: pathlib.Path, output_path: Optional[pathlib.Path] = None) -> pathlib.Path:
        """导出为Markdown格式"""
        if output_path is None:
            output_path = docx_path.with_suffix('.md')
            
        print(f"   📝 导出Markdown: {output_path.name}")
        
        try:
            import pypandoc
            
            # 使用pypandoc转换
            output = pypandoc.convert_file(
                str(docx_path),
                'markdown',
                format='docx',
                extra_args=['--wrap=none', '--extract-media=.']
            )
            
            output_path.write_text(output, encoding='utf-8')
            print(f"   ✅ Markdown导出成功（使用pypandoc）")
            return output_path
            
        except ImportError:
            print(f"   ⚠️  pypandoc未安装，使用简化方案...")
            
            # 简化方案：使用python-docx提取文本
            try:
                from docx import Document
                
                doc = Document(docx_path)
                md_content = []
                
                for para in doc.paragraphs:
                    if para.style.name.startswith('Heading'):
                        level = int(para.style.name[-1]) if para.style.name[-1].isdigit() else 1
                        md_content.append(f"{'#' * level} {para.text}\n")
                    elif para.text.strip():
                        # 处理基本格式
                        text = para.text
                        for run in para.runs:
                            if run.bold:
                                text = text.replace(run.text, f"**{run.text}**")
                            elif run.italic:
                                text = text.replace(run.text, f"*{run.text}*")
                        md_content.append(f"{text}\n")
                
                # 处理表格
                for table in doc.tables:
                    md_content.append("\n")
                    # 表头
                    if table.rows:
                        header_row = table.rows[0]
                        headers = [cell.text.strip() for cell in header_row.cells]
                        md_content.append("| " + " | ".join(headers) + " |")
                        md_content.append("| " + " | ".join(["---"] * len(headers)) + " |")
                        
                        # 数据行
                        for row in table.rows[1:]:
                            cells = [cell.text.strip() for cell in row.cells]
                            md_content.append("| " + " | ".join(cells) + " |")
                    md_content.append("\n")
                
                output_path.write_text('\n'.join(md_content), encoding='utf-8')
                print(f"   ✅ Markdown导出成功（简化方案）")
                return output_path
                
            except Exception as e:
                print(f"   ❌ Markdown导出失败: {e}")
                raise
                
    def export_to_txt(self, docx_path: pathlib.Path, output_path: Optional[pathlib.Path] = None) -> pathlib.Path:
        """导出为纯文本格式"""
        if output_path is None:
            output_path = docx_path.with_suffix('.txt')
            
        print(f"   📃 导出纯文本: {output_path.name}")
        
        try:
            from docx import Document
            
            doc = Document(docx_path)
            text_content = []
            
            # 提取所有段落文本
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content.append(para.text.strip())
                    
            # 提取表格文本
            for table in doc.tables:
                text_content.append("\n[表格]")
                for row in table.rows:
                    row_text = "\t".join(cell.text.strip() for cell in row.cells)
                    text_content.append(row_text)
                text_content.append("")
                
            # 提取页眉页脚
            for section in doc.sections:
                if section.header:
                    header_text = ' '.join(p.text.strip() for p in section.header.paragraphs if p.text.strip())
                    if header_text:
                        text_content.insert(0, f"[页眉] {header_text}")
                        
                if section.footer:
                    footer_text = ' '.join(p.text.strip() for p in section.footer.paragraphs if p.text.strip())
                    if footer_text:
                        text_content.append(f"[页脚] {footer_text}")
                        
            output_path.write_text('\n\n'.join(text_content), encoding='utf-8')
            print(f"   ✅ 纯文本导出成功")
            return output_path
            
        except Exception as e:
            print(f"   ❌ 纯文本导出失败: {e}")
            raise
            
    def export_multiple_formats(self, docx_path: pathlib.Path, formats: List[str]) -> Dict[str, pathlib.Path]:
        """导出多种格式"""
        results = {}
        
        # 格式映射
        format_handlers = {
            'pdf': self.export_to_pdf,
            'html': self.export_to_html,
            'md': self.export_to_markdown,
            'markdown': self.export_to_markdown,
            'txt': self.export_to_txt,
            'text': self.export_to_txt,
            'docx': lambda p: p  # DOCX已存在，直接返回
        }
        
        print(f"\n📤 开始多格式导出...")
        for fmt in formats:
            fmt = fmt.lower().strip()
            if fmt in format_handlers:
                try:
                    output_path = format_handlers[fmt](docx_path)
                    results[fmt] = output_path
                except Exception as e:
                    print(f"   ⚠️  {fmt.upper()}格式导出失败: {e}")
                    results[fmt] = None
            else:
                print(f"   ❌ 不支持的格式: {fmt}")
                results[fmt] = None
                
        # 统计结果
        success_count = sum(1 for v in results.values() if v is not None)
        print(f"\n📊 导出完成: 成功 {success_count}/{len(formats)} 个格式")
        
        return results


# 便捷函数
def export_docx_to_formats(docx_path: str, formats: List[str]) -> Dict[str, Optional[pathlib.Path]]:
    """便捷函数：将DOCX导出为多种格式"""
    exporter = FormatExporter()
    return exporter.export_multiple_formats(pathlib.Path(docx_path), formats)


if __name__ == "__main__":
    # 测试代码
    import sys
    
    if len(sys.argv) < 2:
        print("用法: python exporters.py <docx文件> [格式1,格式2,...]")
        print("支持格式: pdf, html, md, txt")
        sys.exit(1)
        
    docx_file = sys.argv[1]
    formats = sys.argv[2].split(',') if len(sys.argv) > 2 else ['pdf', 'html', 'md', 'txt']
    
    if not pathlib.Path(docx_file).exists():
        print(f"❌ 文件不存在: {docx_file}")
        sys.exit(1)
        
    results = export_docx_to_formats(docx_file, formats)
    
    print("\n✅ 导出结果:")
    for fmt, path in results.items():
        if path:
            print(f"   {fmt.upper()}: {path}")
        else:
            print(f"   {fmt.upper()}: 失败") 