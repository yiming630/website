"""
批量分割DOCX文档脚本
"""
import logging
from pathlib import Path
from typing import List, Tuple, Optional
from docx import Document
from docx.shared import Inches
from tqdm import tqdm
from concurrent.futures import ThreadPoolExecutor, as_completed
from .config import DOCX_RAW_DIR, DOCX_SPLIT_DIR, MAX_WORKERS, SPLIT_MAX_TOKENS
from .gemini_client import GeminiClient

logger = logging.getLogger(__name__)


class DocumentSplitter:
    """文档智能分割器"""
    
    def __init__(self, gemini_client: GeminiClient = None):
        self.gemini_client = gemini_client or GeminiClient()
    
    def extract_text_from_docx(self, docx_path: Path) -> str:
        """从DOCX文件提取纯文本"""
        try:
            doc = Document(docx_path)
            paragraphs = []
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)
            
            # 也提取表格中的文本
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text = cell.text.strip()
                        if text:
                            paragraphs.append(text)
            
            return '\n\n'.join(paragraphs)
            
        except Exception as e:
            logger.error(f"提取文本失败 {docx_path}: {e}")
            raise
    
    def save_segments_to_docx(self, segments: List[str], output_path: Path, 
                             original_docx: Optional[Path] = None):
        """将分割后的文本段保存为DOCX文件"""
        doc = Document()
        
        # 添加标题
        doc.add_heading('智能分割文档', 0)
        
        # 添加元信息
        if original_docx:
            doc.add_paragraph(f'原始文件: {original_docx.name}')
            doc.add_paragraph(f'分割段数: {len(segments)}')
        
        doc.add_page_break()
        
        # 添加各个段落
        for i, segment in enumerate(segments, 1):
            # 段落标题
            doc.add_heading(f'第 {i} 段', level=1)
            
            # 段落内容
            doc.add_paragraph(segment)
            
            # 除了最后一段，其他段后面加分页符
            if i < len(segments):
                doc.add_page_break()
        
        # 保存文档
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        logger.info(f"保存分割文档: {output_path}")
    
    def save_segments_to_txt(self, segments: List[str], output_dir: Path, 
                            base_name: str):
        """将分割后的文本段保存为多个TXT文件"""
        output_dir.mkdir(parents=True, exist_ok=True)
        
        for i, segment in enumerate(segments, 1):
            txt_path = output_dir / f"{base_name}_part{i:03d}.txt"
            with open(txt_path, 'w', encoding='utf-8') as f:
                f.write(segment)
            logger.debug(f"保存文本段: {txt_path}")
    
    def split_single_docx(self, docx_path: Path, output_dir: Path,
                         max_tokens: int = SPLIT_MAX_TOKENS,
                         save_format: str = 'docx') -> Tuple[bool, Path, str]:
        """
        分割单个DOCX文件
        
        Args:
            docx_path: DOCX文件路径
            output_dir: 输出目录
            max_tokens: 每段最大token数
            save_format: 保存格式 ('docx' 或 'txt')
            
        Returns:
            (是否成功, 文件路径, 错误信息)
        """
        try:
            # 提取文本
            text = self.extract_text_from_docx(docx_path)
            
            if not text.strip():
                logger.warning(f"文档为空: {docx_path}")
                return (False, docx_path, "文档内容为空")
            
            # 使用Gemini智能分割
            segments = self.gemini_client.split_text(text, max_tokens)
            
            if not segments:
                logger.warning(f"分割失败: {docx_path}")
                return (False, docx_path, "分割结果为空")
            
            # 保存分割结果
            base_name = docx_path.stem
            
            if save_format == 'docx':
                output_path = output_dir / f"{base_name}_split.docx"
                self.save_segments_to_docx(segments, output_path, docx_path)
                return (True, output_path, "")
            else:
                # 保存为多个TXT文件
                txt_dir = output_dir / base_name
                self.save_segments_to_txt(segments, txt_dir, base_name)
                return (True, txt_dir, "")
                
        except Exception as e:
            error_msg = str(e)
            logger.error(f"分割失败 {docx_path}: {error_msg}")
            return (False, docx_path, error_msg)
    
    def split_batch(self, docx_files: List[Path] = None,
                   input_dir: Path = DOCX_RAW_DIR,
                   output_dir: Path = DOCX_SPLIT_DIR,
                   max_workers: int = MAX_WORKERS,
                   save_format: str = 'docx') -> dict:
        """
        批量分割DOCX文件
        
        Args:
            docx_files: 要分割的DOCX文件列表
            input_dir: 输入目录
            output_dir: 输出目录
            max_workers: 最大并发数
            save_format: 保存格式
            
        Returns:
            分割结果统计
        """
        # 获取DOCX文件列表
        if docx_files is None:
            docx_files = list(input_dir.glob('*.docx'))
        
        if not docx_files:
            logger.warning(f"未找到DOCX文件在: {input_dir}")
            return {
                'total': 0,
                'success': 0,
                'failed': 0,
                'success_files': [],
                'failed_files': []
            }
        
        logger.info(f"准备分割 {len(docx_files)} 个DOCX文件")
        
        # 确保输出目录存在
        output_dir.mkdir(parents=True, exist_ok=True)
        
        # 结果统计
        success_files = []
        failed_files = []
        
        # 使用线程池并发处理
        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            # 提交所有任务
            future_to_docx = {
                executor.submit(
                    self.split_single_docx, 
                    docx_file, 
                    output_dir,
                    SPLIT_MAX_TOKENS,
                    save_format
                ): docx_file
                for docx_file in docx_files
            }
            
            # 使用进度条显示进度
            with tqdm(total=len(docx_files), desc="分割进度") as pbar:
                for future in as_completed(future_to_docx):
                    docx_file = future_to_docx[future]
                    success, file_path, error_msg = future.result()
                    
                    if success:
                        success_files.append(file_path)
                        logger.info(f"✓ 成功: {docx_file.name}")
                    else:
                        failed_files.append((file_path, error_msg))
                        logger.error(f"✗ 失败: {docx_file.name} - {error_msg}")
                    
                    pbar.update(1)
        
        # 统计结果
        result = {
            'total': len(docx_files),
            'success': len(success_files),
            'failed': len(failed_files),
            'success_files': success_files,
            'failed_files': failed_files
        }
        
        logger.info(f"分割完成: 成功 {result['success']}/{result['total']}")
        
        # 如果有失败的文件，记录到错误日志
        if failed_files:
            self._save_error_log(failed_files)
        
        return result
    
    def _save_error_log(self, failed_files: List[Tuple[Path, str]]):
        """保存错误日志"""
        from .config import LOGS_DIR
        from datetime import datetime
        
        error_log_file = LOGS_DIR / 'split_errors.log'
        
        with open(error_log_file, 'a', encoding='utf-8') as f:
            f.write("\n" + "="*50 + "\n")
            f.write(f"分割失败记录 - {datetime.now()}\n")
            for file_path, error_msg in failed_files:
                f.write(f"文件: {file_path}\n错误: {error_msg}\n\n")


def main():
    """主函数"""
    import argparse
    
    parser = argparse.ArgumentParser(description='批量智能分割DOCX文档')
    parser.add_argument('--input', type=Path, default=DOCX_RAW_DIR,
                       help='输入DOCX目录')
    parser.add_argument('--output', type=Path, default=DOCX_SPLIT_DIR,
                       help='输出目录')
    parser.add_argument('--workers', type=int, default=MAX_WORKERS,
                       help='并发工作线程数')
    parser.add_argument('--format', choices=['docx', 'txt'], default='docx',
                       help='输出格式')
    parser.add_argument('--max-tokens', type=int, default=SPLIT_MAX_TOKENS,
                       help='每段最大token数')
    parser.add_argument('--files', nargs='+', type=Path,
                       help='指定要分割的DOCX文件')
    
    args = parser.parse_args()
    
    # 配置日志
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
    )
    
    # 创建分割器并执行
    splitter = DocumentSplitter()
    result = splitter.split_batch(
        docx_files=args.files,
        input_dir=args.input,
        output_dir=args.output,
        max_workers=args.workers,
        save_format=args.format
    )
    
    # 打印结果摘要
    print(f"\n分割完成！")
    print(f"总计: {result['total']} 个文件")
    print(f"成功: {result['success']} 个文件")
    print(f"失败: {result['failed']} 个文件")
    
    if result['failed'] > 0:
        print(f"\n失败文件详情请查看: logs/split_errors.log")


if __name__ == '__main__':
    main() 